import tkinter as tk
from tkinter import filedialog
from docx import Document
import re
import datetime
from docxlatex import Document as Form

WORK_DICT = {
    "1ф": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии однофазного",
    "3ф ПР": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного непосредственного (прямого) включения",
    "3ф ПК": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного трансформаторного включения"
}
EQUIP_DICCT = {
    "1ф": "Счетчик электрической энергии однофазный, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM",
    "3ф ПР": "Счетчик электрической энергии трехфазный непосредственного (прямого) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.749.OG.QYUVLFZ NB02 SPds",
    "3ф ПК": "Счетчик электрической энергии трехфазный трансформаторного (полукосвенного) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.543.OAG.SYUVLFZ NB02 SPds"
}


def find_info_in_doc(doc, path):
    info = {
        "номер заявки": 0,
        "дата обращения": "",
        "таблица работ": [],
        "суммарная стоимость": 0,
        "дата выполнения": "",
        "адреса объектов": [],
        "режим": 1
    }

    # === Извлечение текста из параграфов ===
    full_text = "\n".join([p.text for p in doc.paragraphs])

    # === Поиск текстовой информации ===

    # Номер заявки
    match = re.search(r'Заявка\s*на\s*работы\s*№\s*(\d+)', full_text, re.IGNORECASE)
    if match:
        info["номер заявки"] = match.group(0).split("№")[1].strip()

    # Дата обращения
    dat = ""
    Temp = Form(path)
    equations = Temp.get_text()
    eq_l = equations.replace("\n", '**').replace("\t", '**')
    eq_l = eq_l.split("**")
    for el in eq_l:
        if not el.strip().replace(' ', ''):
            eq_l.remove(el)
    for el in eq_l:
        if "овосибирск" in el.strip().lower():
            dat = eq_l[eq_l.index(el)+1].strip()
            break
    info["дата обращения"] = dat

    # Дата выполнения работ
    match = re.search(r'Дата\s*выполнения\s*работ:\s*([^\n\r]+)', full_text, re.IGNORECASE)
    if match:
        info["дата выполнения"] = match.group(0).strip()

    # Таблица работ: ищем таблицу с заголовками работ
    for table in doc.tables:
        if table.rows:
            first_row_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            first_row_text = " ".join(first_row_cells)

            if any([x in first_row_text for x in ["наименование", "работ", "стоимость"]]):
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data) and not all(cell == "" for cell in row_data):
                        if len(row_data) > 0 and not re.match(r'^(Работы|Оборудование)[:\s]*$', row_data[0]):
                            info["таблица работ"].append(row_data)

                # Ищем итоговую стоимость
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells])
                    if "итоговая стоимость" in row_text.lower():
                        # Берем последнюю ячейку
                        last_cell = row.cells[-1].text.strip()
                        if last_cell:
                            info["суммарная стоимость"] = last_cell
                break

    # Таблица адресов объектов
    for table in doc.tables:
        if table.rows:
            first_row_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            first_row_text = " ".join(first_row_cells)

            if "адрес объекта" in first_row_text or ("№п/п" in first_row_text and "адрес" in first_row_text):
                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data) and not all(cell == "" for cell in row_data):
                        info["адреса объектов"].append(row_data)
                break

    prep = []
    flag = False
    if not info["адреса объектов"]:
        full_l = full_text.split("\n")    #
        for i in range(len(full_l)):
            el = full_l[i]
            if "работ:" in el.strip():
                break
            elif flag and any([x in el.lower() for x in ["г.", "п.", "о.", "1.", "с.",
                                                         "город", "посел", "област", "сел", "пгт"]]):
                prep += [el.split()]
            elif "адрес" in el.strip().lower() and "объект" in el.strip().lower():
                flag = True
                if any([x in el.lower() for x in ["г.", "п.", "о.", "1.", "с.",
                                                  "город", "посел", "област", "сел", "пгт"]]):
                    prep += [el.split(":")[1].strip().split()]

        tab_data = []    #
        ind_num = 0
        pu_num = int(info["таблица работ"][1][3]) // len(prep)
        typ_py = ""
        adress = ""
        addit = ""

        for el in prep:

            table2 = info["таблица работ"][1:-2]  # Убираем первые 2 и последние 2 элемента
            if table2[0][1].strip() == "":
                table2.remove(table2[0])
            mid_point = len(table2) // 2
            table2_work = table2[:mid_point]  # Данные работ

            type_counter = {"1ф": 0, "3ф ПР": 0, "3ф ПК": 0}
            for row in table2_work:
                if "однофазн" in row[1]:
                    type_counter["1ф"] += 1
                if ("трехфазн" in row[1]) and any([w in row[1] for w in ["непосредств", "прям"]]):
                    type_counter["3ф ПР"] += 1
                if ("трехфазн" in row[1]) and any([w in row[1] for w in ["полукосвенн", "трансформаторн"]]):
                    type_counter["3ф ПК"] += 1

            for sub in el:
                if "кв." == sub.strip().lower():
                    ind = el.index(sub)
                    adress = " ".join(el[:ind + 2])
                    if len(el) >= ind + 3:
                        addit = " ".join(el[ind + 2:])
            if not adress:
                adress = " ".join(el)

            for typ in list(type_counter.keys()):
                if type_counter[typ]:
                    ind_num += 1
                    row = table2_work.pop()
                    typ_py = typ

                    pu_num = row[3]

                    tab_row = [str(ind_num), str(pu_num), typ_py, adress, addit]
                    tab_data += [tab_row]



        info["адреса объектов"] = tab_data
        info["режим"] = 0

    print(info["адреса объектов"])
    return info


def open_file(path):
    file_path = path
    if not file_path:
        return

    # Дата обращения - ищем дату в первых строках

    doc = Document(file_path)
    info = find_info_in_doc(doc)


    # Вывод информации
    print("\n=== Извлечённая информация ===")
    print("Номер заявки:", info["номер заявки"])
    print("Дата обращения:", info["дата обращения"])
    print("Дата выполнения:", info["дата выполнения"])
    print("Суммарная стоимость:", info["суммарная стоимость"])
    print("Таблица работ:")
    for row in info["таблица работ"]:
        print("  ", row)
    print("Адреса объектов:")
    for row in info["адреса объектов"]:
        print("  ", row)

    return info


