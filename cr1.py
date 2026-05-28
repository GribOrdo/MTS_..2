from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words

WORK_DICT = {
    "1ф": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии однофазного",
    "3ф ПР": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного непосредственного (прямого) включения",
    "3ф ПК": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного трансформаторного включения"
}
EQUIP_DICT = {
    "1ф": "Счетчик электрической энергии однофазный, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM",
    "3ф ПР": "Счетчик электрической энергии трехфазный непосредственного (прямого) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.749.OG.QYUVLFZ NB02 SPds",
    "3ф ПК": "Счетчик электрической энергии трехфазный трансформаторного (полукосвенного) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.543.OAG.SYUVLFZ NB02 SPds"
}
PR_WK_DICT = {
    "1ф": "1250,00",
    "3ф ПР": "1250,00",
    "3ф ПК": "1250,00"
}
PR_EQ_DICT = {
    "1ф": "0,00",
    "3ф ПР": "0,00",
    "3ф ПК": "0,00"
}


def price_to_show(pr):
    price_show = format(float(pr), ",.2f").replace(",", " ")
    return price_show


def sum_to_words(amount):
    rub = int(amount)
    kop = int(round((amount - rub) * 100))

    rub_text = num2words(rub, lang='ru')
    kop_text = num2words(kop, lang='ru')

    return f"{rub_text} руб. {kop_text} коп."


def clean_price_string(price_str):
    """Очистка строки с ценой для конвертации в float"""
    price_str = str(price_str)
    price_str = price_str.replace('\xa0', ' ')  # Заменяем неразрывные пробелы
    price_str = price_str.replace(' ', '')  # Удаляем все пробелы
    price_str = price_str.replace(',', '.')  # Заменяем запятую на точку
    if price_str.strip() == "":
        return 0.00
    return float(price_str)


def add_formatted_paragraph(doc, text, bold=False, size=11, alignment=None, font_name='Times New Roman'):
    """Добавление форматированного параграфа"""
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    return para


def create_application_doc(num, date, table1_rows, table1_data, table2_data, date_work, mode):
    """Создание документа заявки"""

    doc = Document()

    # Настройка стилей документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Заголовок
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f'Заявка на работы №{num}')
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Times New Roman'

    # Местоположение и дата
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = header_para.add_run('г. Новосибирск')
    run1.font.size = Pt(11)
    run1.font.name = 'Times New Roman'

    # Добавляем табуляцию для даты справа
    run2 = header_para.add_run(f'\t\t\t\t\t\t\t\t\t{date}')
    run2.font.size = Pt(11)

    run3 = header_para.add_run('\n\t1. Список объектов:')
    run3.font.size = Pt(12)
    run3.font.name = 'Times New Roman'

    # Создание таблицы объектов
    table1 = doc.add_table(rows=table1_rows, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки таблицы
    headers = ['порядковый номер', 'кол-во ПУ', 'тип ПУ', 'Адрес объекта', 'Примечания']
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные таблицы
    data = table1_data

    # Заполнение таблицы данными
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table1.rows[row_idx].cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_data))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if col_idx <= 2:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Установка ширины столбцов
    widths = [Cm(1.24), Cm(1.41), Cm(1.59), Cm(6.1), Cm(7.58)]
    for row in table1.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()

    # Подсчет типов счетчиков
    type_counter = {"1ф": 0, "3ф ПР": 0, "3ф ПК": 0}
    for row in table1_data:
        if "1ф" in row[2]:
            type_counter["1ф"] += 1
        if "ПР" in row[2]:
            type_counter["3ф ПР"] += 1
        if "ПК" in row[2]:
            type_counter["3ф ПК"] += 1

    # Извлекаем данные из table2_data
    table2 = table2_data[1:-2]  # Убираем первые 2 и последние 2 элемента
    if table2[0][1].strip() == "":
        table2.remove(table2[0])
    mid_point = len(table2) // 2
    table2_work = table2[:mid_point]  # Данные работ
    table2_equip = table2[mid_point:]  # Данные оборудования

    for work in table2_work:
        if "однофазн" in work[1]:
            WORK_DICT["1ф"] = work[1].strip()
        if ("трехфазн" in work[1]) and any([w in work[1] for w in ["непосредств", "прям"]]):
            WORK_DICT["3ф ПР"] = work[1].strip()
        if ("трехфазн" in work[1]) and any([w in work[1] for w in ["полукосвенн", "трансформаторн"]]):
            WORK_DICT["3ф ПК"] = work[1].strip()

    for eq in table2_equip:
        if "однофазн" in eq[1]:
            EQUIP_DICT["1ф"] = eq[1].strip()
        if ("трехфазн" in eq[1]) and any([w in eq[1] for w in ["непосредств", "прям"]]):
            EQUIP_DICT["3ф ПР"] = eq[1].strip()
        if ("трехфазн" in eq[1]) and any([w in eq[1] for w in ["полукосвенн", "трансформаторн"]]):
            EQUIP_DICT["3ф ПК"] = eq[1].strip()

    # Раздел 2: Перечень работ
    add_formatted_paragraph(doc, '\t2. Перечень работ:', size=11)

    # Создаем таблицу перечня работ (начинаем с 2 строк: заголовок + "Работы:")
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки таблицы 2
    headers2 = ['№ п/п', 'Наименование работ', 'Единица измерения', 'Объем выполняемых работ',
                'Цена работ за единицу, руб. без НДС¹', 'Стоимость работ, руб. без НДС¹']

    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = False
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Объединение ячеек для "Работы:"
    cell_work = table2.rows[1].cells[0]
    merged_cell = cell_work.merge(table2.rows[1].cells[5])
    merged_cell.text = ''
    para = merged_cell.paragraphs[0]
    run = para.add_run('Работы:')
    run.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Добавляем строки с данными работ
    table2_p = 1
    total_cost = []

    for typ in WORK_DICT.keys():
        if mode == 1:
            if type_counter[typ] > 0:
                if table2_work:
                    row_data = table2_work.pop(0)
                    price_per_unit = clean_price_string(PR_WK_DICT[typ])

                    total_cost.append(type_counter[typ] * price_per_unit)

                    work_values = [
                        str(table2_p),
                        WORK_DICT[typ],
                        str(row_data[2]),
                        str(type_counter[typ]),
                        price_to_show(price_per_unit),
                        price_to_show(type_counter[typ] * price_per_unit)
                    ]

                    row = table2.add_row()
                    for col_idx, cell_data in enumerate(work_values):
                        cell = row.cells[col_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        run = para.add_run(str(cell_data))
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        if col_idx in [0, 2, 3, 4, 5]:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    table2_p += 1
        else:
            if table2_work:
                row_data = table2_work.pop(0)
                price_per_unit = clean_price_string(PR_WK_DICT[typ])
                total_cost.append(int(row_data[3]) * price_per_unit)

                work_values = [
                    str(table2_p),
                    WORK_DICT[typ],
                    str(row_data[2]),
                    str(row_data[3]),
                    price_to_show(price_per_unit),
                    price_to_show(int(row_data[3]) * price_per_unit)
                ]

                row = table2.add_row()
                for col_idx, cell_data in enumerate(work_values):
                    cell = row.cells[col_idx]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_data))
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if col_idx in [0, 2, 3, 4, 5]:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table2_p += 1

    # Добавляем строку "Оборудование:"
    row_eq_header = table2.add_row()
    merged_cell2 = row_eq_header.cells[0].merge(row_eq_header.cells[5])
    merged_cell2.text = ''
    para = merged_cell2.paragraphs[0]
    run = para.add_run('Оборудование:')
    run.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Добавляем строки с данными оборудования
    for typ in EQUIP_DICT.keys():
        if mode == 1:
            if type_counter[typ] > 0:
                if table2_equip:
                    row_data = table2_equip.pop(0)
                    price_per_unit = 0
                    total_cost.append(type_counter[typ] * price_per_unit)

                    equip_values = [
                        str(table2_p),
                        EQUIP_DICT[typ],
                        str(row_data[2]),
                        str(type_counter[typ]),
                        price_to_show(price_per_unit),
                        price_to_show(type_counter[typ] * price_per_unit)
                    ]

                    row = table2.add_row()
                    for col_idx, cell_data in enumerate(equip_values):
                        cell = row.cells[col_idx]
                        cell.text = ''
                        para = cell.paragraphs[0]
                        run = para.add_run(str(cell_data))
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        if col_idx in [0, 2, 3, 4, 5]:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    table2_p += 1
        else:
            if table2_equip:
                row_data = table2_equip.pop(0)
                price_per_unit = 0
                total_cost.append(type_counter[typ] * price_per_unit)

                equip_values = [
                    str(table2_p),
                    EQUIP_DICT[typ],
                    str(row_data[2]),
                    str(row_data[3]),
                    price_to_show(price_per_unit),
                    price_to_show(int(row_data[3]) * price_per_unit)
                ]

                row = table2.add_row()
                for col_idx, cell_data in enumerate(equip_values):
                    cell = row.cells[col_idx]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_data))
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if col_idx in [0, 2, 3, 4, 5]:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table2_p += 1

    total_sum = sum(total_cost)

    # Итоговая строка
    total_row = table2.add_row()
    merged_total = total_row.cells[0].merge(total_row.cells[4])
    merged_total.text = ''
    para = merged_total.paragraphs[0]
    run = para.add_run('Итоговая стоимость работ по Заявке, руб. без НДС¹')
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    total_cell = total_row.cells[5]
    total_cell.text = ''
    para = total_cell.paragraphs[0]
    run = para.add_run(price_to_show(total_sum))
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Установка ширины столбцов
    widths = [Cm(0.99), Cm(8), Cm(1.5), Cm(2), Cm(2.75), Cm(2.29)]
    for row in table2.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Сноска
    footnote = doc.add_paragraph()
    footnote_run = footnote.add_run(
        '¹Размер НДС определяется по ставке, установленной п. 3 ст. 164 НК РФ.')
    footnote_run.font.size = Pt(11)
    footnote_run.font.name = 'Times New Roman'

    # Пункт 3
    para3 = doc.add_paragraph()
    run3 = para3.add_run(
        '3. Требования к составу материалов и оборудования: Договором предусмотрено давальческие материалы/оборудование в составе (Заполняется при наличии давальческих материалов/оборудования):')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    # Таблица давальческих материалов
    table3 = doc.add_table(rows=2, cols=4)
    table3.style = 'Table Grid'

    materials_headers = ['№ п/п', 'Наименование давальческих материалов/оборудования', 'Единица измерения',
                         'Количество']
    for i, header in enumerate(materials_headers):
        cell = table3.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.bold = False
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Строка с прочерками
    dash_data = ['1', '-', '-', '-']
    for i, data in enumerate(dash_data):
        cell = table3.rows[1].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(data)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    widths = [Cm(1.43), Cm(8.9), Cm(3.33), Cm(3.37)]
    for row in table3.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Раздел 4
    total_nds = total_sum * 1.22
    nds_amount = total_nds - total_sum
    add_formatted_paragraph(doc,
                            f'4. Стоимость работ, выполняемых по настоящей Заявке определена в соответствии с условиями Договора и составляет {price_to_show(total_nds)} ({sum_to_words(total_nds)}), в том числе НДС 22% {price_to_show(nds_amount)} ({sum_to_words(nds_amount)}).',
                            bold=False, size=11)

    # Раздел 5
    add_formatted_paragraph(doc,
                            '5. Окончательная стоимость работ определяется исходя из фактических объемов, зафиксированных в Акте о приемке выполненных работ, и не может превышать стоимость работ, указанных в п.4 настоящей Заявки.',
                            bold=False, size=11)

    # Раздел 6
    add_formatted_paragraph(doc,
                            '6. В случае, если фактическая стоимость работ оказалась меньше суммы, указанной в п. 4. настоящей Заявки, разница между суммой, указанной в п. 4 настоящей Заявки и стоимостью фактически выполненных работ остается у Заказчика.',
                            bold=False, size=11)

    # Раздел 7
    add_formatted_paragraph(doc,
                            f'7. {date_work}',
                            bold=False, size=11)

    # Подписи
    add_formatted_paragraph(doc, 'Заказчик', size=11)
    add_formatted_paragraph(doc, 'Должность:', size=11)
    add_formatted_paragraph(doc, '__________________ / ___________ / М.П.', size=10)

    # Сохранение документа
    output_path = 'out1.docx'
    doc.save(output_path)


    return output_path
