from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words

WORK_DICT = {
    "1ф": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии однофазного",
    "3ф ПР": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного непосредственного (прямого) включения",
    "3ф ПК": "Электромонтажные и пусконаладочные работы по подключению счетчика электрической энергии трехфазного трансформаторного включения"
}
EQUIP_DICT = {
    "1ф": "Счетчик электрической энергии однофазный, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM",
    "3ф ПР": "Счетчик электрической энергии трехфазный непосредственного (прямого) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.749.OG.QYUVLFZ NB02 SPds",
    "3ф ПК": "Счетчик электрической энергии трехфазный трансформаторного (полукосвенного) включения, соответствующий требованиям ПП РФ № 890 от 19.06.2020 г., NBIOT/GSM_CE307 R34.543.OAG.SYUVLFZ NB02 SPds"
}
PR_WK_DICT = {
    "1ф": "1764,71",
    "3ф ПР": "1764,71",
    "3ф ПК": "1764,71"
}
PR_EQ_DICT = {
    "1ф": "7434,12",
    "3ф ПР": "12352,94",
    "3ф ПК": "11176,48"
}

def clean_price_string(price_str):
    """Очистка строки с ценой для конвертации в float"""
    price_str = str(price_str)
    price_str = price_str.replace('\xa0', ' ')  # Заменяем неразрывные пробелы
    price_str = price_str.replace(' ', '')  # Удаляем все пробелы
    price_str = price_str.replace(',', '.')  # Заменяем запятую на точку
    if price_str.strip() == "":
        return 0.00
    return float(price_str)

def price_to_show(pr):
    price_show = format(float(pr), ",.2f").replace(",", " ")
    return price_show


def sum_to_words(amount):
    rub = int(amount)
    kop = int(round((amount - rub) * 100))

    rub_text = num2words(rub, lang='ru')
    kop_text = num2words(kop, lang='ru')

    return f"{rub_text} руб. {kop_text} коп."


def create_application_doc2(num, date1, date2, table1_rows, table1_data, table2_data, mode):
    # Извлекаем данные из table2_data (без заголовков и итогов)
    # Предполагаемая структура: ['Работы:', данные_работ, 'Оборудование:', данные_оборудования, 'ИТОГО:', 'НДС:']
    table2 = table2_data[1:-2]  # Убираем первые 2 и последние 2 элемента
    if table2[0][1].strip() == "":
        table2.remove(table2[0])
    mid_point = len(table2) // 2
    table2_work = table2[:mid_point]
    table2_equip = table2[mid_point:]

    for work in table2_work:
        if "однофазн" in work[1]:
            WORK_DICT["1ф"] = work[1].strip()
        if ("трехфазн" in work[1]) and any([w in work[1] for w in ["непосредств", "прям"]]):
            WORK_DICT["3ф ПР"] = work[1].strip()
        if ("трехфазн" in work[1]) and any([w in work[1] for w in ["полукосвенн", "трансформаторн"]]):
            WORK_DICT["3ф ПК"] = work[1].strip()

    for eq in table2_equip:
        if "однофазн" in eq[1]:
            EQUIP_DICT["1ф"] = eq[1].strip()
        if ("трехфазн" in eq[1]) and any([w in eq[1] for w in ["непосредств", "прям"]]):
            EQUIP_DICT["3ф ПР"] = eq[1].strip()
        if ("трехфазн" in eq[1]) and any([w in eq[1] for w in ["полукосвенн", "трансформаторн"]]):
            EQUIP_DICT["3ф ПК"] = eq[1].strip()


    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Установка узких полей
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)

    # Заголовок акта
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f'Акт о приемке выполненных работ № {num}')
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run.font.name = 'Times New Roman'

    # Таблица для даты и города
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cells = header_table.rows[0].cells
    cells[0].text = 'г. Новосибирск'
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cells[1].text = f'{date1}'
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cells[0].width = Cm(9)
    cells[1].width = Cm(9)

    # Убираем границы таблицы
    for cell in header_table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    main_t1 = 'Акционерное общество «Новосибирскэнергосбыт» (АО «Новосибирскэнергосбыт»), '
    main_t2 = 'именуемое в дальнейшем '
    main_t3 = '«Заказчик»'
    main_t4 = ', в лице Заместителя генерального директора — Технического директора Михайлишина Александра Юрьевича, действующего на основании Доверенности № 24/2024 от 02.02.2024 г., с одной стороны, и '
    main_t5 = 'Публичное акционерное общество «Мобильные ТелеСистемы» (ПАО «МТС») '
    main_t6 = 'именуемое в дальнейшем'
    main_t7 = ' «Подрядчик»'
    main_t8 = (', в лице Директора Департамента по работе с '
               'корпоративными клиентами Ватуля Елены Николаевны, действующего на основании '
               'Доверенности № 77/509-н/77-2024-3-661 от 19.09.2024 г., другой стороны, '
               'составили настоящий Акт к договору подряда №39278 от 28.02.2025 г. в том, что:')

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.add_run(main_t1).bold = True
    para.add_run(main_t2)
    para.add_run(main_t3).bold = True
    para.add_run(main_t4)
    para.add_run(main_t5).bold = True
    para.add_run(main_t6)
    para.add_run(main_t7).bold = True
    para.add_run(main_t8)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Пункт 1
    para1 = doc.add_paragraph()
    para1.paragraph_format.first_line_indent = Cm(1.25)
    run1 = para1.add_run(
        f'1. В соответствии с заключенным между Подрядчиком и Заказчиком договором, '
        f'Подрядчик выполнил, а Заказчик принял следующие работы по Заявке на работы '
        f'№ {num} от {date2}'
    )
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)

    # Пункт 2 - таблица объектов
    para2 = doc.add_paragraph()
    para2.paragraph_format.first_line_indent = Cm(1.25)
    run2 = para2.add_run('2. Объекты по адресу(ам):')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

    # Создание таблицы объектов
    if mode==1:
        # Создание таблицы объектов
        objects_data = table1_data
        table = doc.add_table(rows=table1_rows, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки таблицы
        headers = ['порядковый номер', 'кол-во ПУ', 'тип ПУ', 'Адрес объекта', 'Примечания']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

        # Заполнение данных
        for i, row_data in enumerate(objects_data):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

        # Настройка ширины колонок
        widths = [Cm(1.24), Cm(1.41), Cm(1.59), Cm(6.1), Cm(7.58)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        doc.add_paragraph()

    else:
        objects_data = table1_data
        # headers = ['порядковый номер', 'кол-во ПУ', 'тип ПУ', 'Адрес объекта', 'Примечания']
        ppp = 1
        adresses = set([x[3] for x in objects_data])
        for adr in adresses:
            para_ad = doc.add_paragraph()
            para_ad.paragraph_format.first_line_indent = Cm(2)
            dop = "; ".join(set([x[4] for x in objects_data if adr in x]))
            ad_out = ". ".join([par.capitalize() for par in adr.split(". ")]).strip(" ").strip(".").strip(";").strip(",")
            if ad_out.startswith("1."):
                ad_out = ad_out[4:].lstrip(" ")
            pu_n = [x[1] for x in objects_data if adr in x]
            pu_t = [x[2] for x in objects_data if adr in x]
            pu_p = []
            for i in range(sum([int(adr == x[3]) for x in objects_data])):
                pu_p.append(f'{pu_t[i]} ПУ - {pu_n[i]}')

            run_ad = para_ad.add_run(f'2.{ppp}. {ad_out}. Установленные ПУ: {", ".join(pu_p)}. {dop.strip(" ").strip(".").strip(";").strip(",")}')
            run_ad.font.name = 'Times New Roman'
            run_ad.font.size = Pt(12)
            ppp += 1


    type_counter = {"1ф": 0, "3ф ПР": 0, "3ф ПК": 0}
    for row in table1_data:
        if "1ф" in row[2]:
            type_counter["1ф"] += int(row[1]) if row[1] else 1
        if "ПР" in row[2]:
            type_counter["3ф ПР"] += int(row[1]) if row[1] else 1
        if "ПК" in row[2]:
            type_counter["3ф ПК"] += int(row[1]) if row[1] else 1

    # Пункт 3 - Таблица работ
    para3 = doc.add_paragraph()
    para3.paragraph_format.first_line_indent = Cm(1.25)
    run3 = para3.add_run('3. Перечень выполненных работ:')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # Создаем таблицу перечня работ (начинаем с 2 строк: заголовок + "Работы:")
    table2 = doc.add_table(rows=2, cols=6)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER


    # Заголовки (строка 0)
    work_headers = ['№ п/п', 'Наименование работ', 'Единица измерения',
                    'Объем выполняемых работ', 'Цена работ за единицу, руб. с НДС',
                    'Стоимость работ, руб. с НДС']

    for i, header in enumerate(work_headers):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)

    # Заголовок "Работы:" (строка 1)
    cell_work = table2.rows[1].cells[0]
    merged_cell = cell_work.merge(table2.rows[1].cells[5])
    merged_cell.text = ''
    para = merged_cell.paragraphs[0]
    run = para.add_run('Работы:')
    run.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    table2_p = 1
    total_cost = []

    # Заполняем данные работ
    for typ in WORK_DICT.keys():
        #   if mode == 1:
        if type_counter[typ] > 0:
            if table2_work:  # Проверяем, есть ли данные
                row_data = table2_work.pop(0)
                price_per_unit = clean_price_string(PR_WK_DICT[typ])

                total_cost.append(type_counter[typ] * price_per_unit)

                work_values = [
                    str(table2_p),
                    WORK_DICT[typ],
                    str(row_data[2]),
                    str(type_counter[typ]),
                    price_to_show(price_per_unit),
                    price_to_show(type_counter[typ] * price_per_unit)
                ]

                row = table2.add_row()
                for col_idx, cell_data in enumerate(work_values):
                    cell = row.cells[col_idx]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_data))
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if col_idx in [0, 2, 3, 4, 5]:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table2_p += 1


    # Добавляем строку "Оборудование:"
    row_eq_header = table2.add_row()
    merged_cell2 = row_eq_header.cells[0].merge(row_eq_header.cells[5])
    merged_cell2.text = ''
    para = merged_cell2.paragraphs[0]
    run = para.add_run('Оборудование:')
    run.bold = False
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Заполняем данные оборудования
    for typ in EQUIP_DICT.keys():
        #if mode == 1:
        if type_counter[typ] > 0:
            if table2_equip:
                row_data = table2_equip.pop(0)
                price_per_unit = clean_price_string(PR_EQ_DICT[typ])
                total_cost.append(type_counter[typ] * price_per_unit)

                equip_values = [
                    str(table2_p),
                    EQUIP_DICT[typ],
                    str(row_data[2]),
                    str(type_counter[typ]),
                    price_to_show(price_per_unit),
                    price_to_show(type_counter[typ] * price_per_unit)
                ]

                row = table2.add_row()
                for col_idx, cell_data in enumerate(equip_values):
                    cell = row.cells[col_idx]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_data))
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if col_idx in [0, 2, 3, 4, 5]:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table2_p += 1

    # Итоговая строка
    nds_amount = (sum(total_cost)/1.22 - sum(total_cost)) * (-1)
    total_sum = sum(total_cost)

    # Итоговая строка
    total_row = table2.add_row()
    merged_total = total_row.cells[0].merge(total_row.cells[4])
    merged_total.text = ''
    para = merged_total.paragraphs[0]
    run = para.add_run('ИТОГО')
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    total_cell = total_row.cells[5]
    total_cell.text = ''
    para = total_cell.paragraphs[0]
    run = para.add_run(price_to_show(total_sum))
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Строка НДС
    nds_row = table2.add_row()
    merged_nds = nds_row.cells[0].merge(nds_row.cells[4])
    merged_nds.text = ''
    para = merged_nds.paragraphs[0]
    run = para.add_run('в том числе НДС 22%')
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    nds_cell = nds_row.cells[5]
    nds_cell.text = ''
    para = nds_cell.paragraphs[0]
    run = para.add_run(price_to_show(nds_amount))
    run.bold = False
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Настройка ширины колонок
    work_widths = [Cm(0.99), Cm(6), Cm(1.5), Cm(2), Cm(2.75), Cm(2.29)]
    for row in table2.rows:
        if len(row.cells) == 6:
            for idx, width in enumerate(work_widths):
                row.cells[idx].width = width

    # Сноска
    footnote = doc.add_paragraph()
    footnote_run = footnote.add_run(
        '¹Размер НДС определяется по ставке, установленной п. 3 ст. 164 НК РФ.')
    footnote_run.font.size = Pt(11)
    footnote_run.font.name = 'Times New Roman'

    # Пункт 3
    para3 = doc.add_paragraph()
    para3.paragraph_format.first_line_indent = Cm(1.25)
    run3 = para3.add_run('3. Требования к составу материалов и оборудования: Договором предусмотрено давальческие материалы/оборудование в составе (Заполняется при наличии давальческих материалов/оборудования):')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)

    # Пункт 4
    para4 = doc.add_paragraph()
    para4.paragraph_format.first_line_indent = Cm(1.25)
    run4 = para4.add_run(
        f'4. Всего выполнено работ на сумму: {price_to_show(sum(total_cost))} руб. '
        f'({sum_to_words(sum(total_cost))}), '
        f'в том числе НДС 22% - {price_to_show(nds_amount)} руб. '
        f'({sum_to_words(nds_amount)}).'
    )
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)

    # Пункт 5
    para5 = doc.add_paragraph()
    para5.paragraph_format.first_line_indent = Cm(1.25)
    run5 = para5.add_run(
        '5. Вышеперечисленные работы выполнены полностью и в срок. '
        'Заказчик претензий по объему, качеству и срокам выполненных работ не имеет.'
    )
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(12)

    # Пункт 6
    para6 = doc.add_paragraph()
    para6.paragraph_format.first_line_indent = Cm(1.25)
    run6 = para6.add_run(
        '6. Настоящий Акт составлен в двух идентичных экземплярах, имеющих равную '
        'юридическую силу, по одному для каждой стороны.'
    )
    run6.font.name = 'Times New Roman'
    run6.font.size = Pt(12)

    doc.add_paragraph()

    # Таблица подписей
    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.style = 'Table Grid'

    for row in sign_table.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(9)

    # Заголовки
    cell_left = sign_table.rows[0].cells[0]
    cell_left.text = ''
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Подрядчик:')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    cell_right = sign_table.rows[0].cells[1]
    cell_right.text = ''
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Заказчик:')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Подписи
    cell_left_sign = sign_table.rows[1].cells[0]
    cell_left_sign.text = ''
    for text_line in ['ПАО «МТС»\n', '', 'Ведущий специалист группы выпуска региональных счетов и корреспонденции\n', '',
                      '____________________/ Ложкин М.М. /', '', 'м.п.']:
        if text_line:
            p = cell_left_sign.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if text_line == "ПАО «МТС»\n":
                run.bold = True

    cell_right_sign = sign_table.rows[1].cells[1]
    cell_right_sign.text = ''
    for text_line in ['АО «Новосибирскэнергосбыт»\n', '',
                      'Заместитель генерального директора -',
                      'Технический директор\n', '',
                      '____________________/Михайлишин А.Ю./', '', 'м.п.']:
        if text_line:
            p = cell_right_sign.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if text_line == "АО «Новосибирскэнергосбыт»\n":
                run.bold = True

    output_path = 'out2.docx'
    doc.save(output_path)
    print(f"Документ сохранен как: {output_path}")

    return output_path